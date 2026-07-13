import os
import io
import re
import pdfplumber
from docx import Document
from docx.shared import Inches
import streamlit as st

# Set up the web page title, icon, and layout profile
st.set_page_config(
    page_title="රුහුණු භාණ්ඩ ලැයිස්තුව", 
    page_icon="📋", 
    layout="centered",
    initial_sidebar_state="collapsed"
)

# Custom CSS for soft light blue palette, black fonts, and light green table hovers
st.markdown("""
    <style>
    /* Main background and global font color */
    .stApp {
        background-color: #F0F4F8;
    }
    
    html, body, [data-testid="stWidgetLabel"], p, div, h1, h2, h3, span {
        color: #000000 !important;
        font-family: 'Segoe UI', Helvetica, Arial, sans-serif;
    }
    
    /* Header title styling */
    .main-title {
        color: #0D47A1;
        font-size: 2.4rem;
        font-weight: 600;
        text-align: center;
        margin-bottom: 0.5rem;
    }
    
    .sub-title {
        color: #333333;
        font-size: 1.1rem;
        text-align: center;
        margin-bottom: 2.5rem;
    }
    
    /* Light Blue elegant card styling for data preview */
    .preview-card {
        background-color: #FFFFFF;
        border-radius: 8px;
        padding: 1.5rem;
        box-shadow: 0 4px 6px rgba(0,0,0,0.02), 0 1px 3px rgba(0,0,0,0.05);
        border: 2px solid #BBDEFB;
        margin-top: 1.5rem;
    }
    
    /* Force Dataframe tables to use light green highlight color during hovers instead of black */
    [data-testid="stTable"] tr:hover, 
    [data-testid="stDataFrame"] tr:hover,
    div[data-role="grid"] div[role="row"]:hover {
        background-color: #E8F5E9 !important;
    }
    
    /* Ensure internal canvas data viewer matches light cell styling accents */
    .glideDataEditor-canvas {
        --bg-color-hover: #E8F5E9 !important;
        --accent-color: #C8E6C9 !important;
    }
    
    /* Soft Blue Step indicators style */
    .step-container {
        background-color: #E3F2FD;
        padding: 1rem;
        border-left: 5px solid #2196F3;
        border-radius: 4px;
        margin-bottom: 1.5rem;
        color: #000000;
        font-weight: 500;
    }
    
    /* File Uploader custom styling adjustments */
    [data-testid="stFileUploaderDropzone"] {
        background-color: #FFFFFF;
        border: 2px dashed #90CAF9 !important;
        border-radius: 8px;
    }
    
    /* Custom style tables to maintain black text grid lines */
    table {
        color: #000000 !important;
    }
    
    /* Hide default Streamlit decorations */
    footer {visibility: hidden;}
    #MainMenu {visibility: hidden;}
    </style>
""", unsafe_allow_html=True)

# Hardcoded Clean Sinhala Product Mapping Dictionary
PRODUCT_MAPPING = {
    "roasted chilli powder 50 g": "බැදපු මිරිස් කුඩු 50 g",
    "roasted chilli powder 50gm": "බැදපු මිරිස් කුඩු 50 g",
    "turmeric powder 25 g": "කහ කුඩු 25 g",
    "turmeric powder 25gm": "කහ කුඩු 25 g",
    "pepper powder 250 g": "ගම්මිරිස් කුඩු 250 g",
    "pepper powder 250gm": "ගම්මිරිස් කුඩු 250 g",
    "pepper powder 25 g": "ගම්මිරිස් කුඩු 25 g",
    "pepper powder 25gm": "ගම්මිරිස් කුඩු 25 g",
    "plums 50 g": "වියළි මිදි / ප්ලම්ස් 50 g",
    "plums 50gm": "වියළි මිදි / ප්ලම්ස් 50 g",
    "meat curry mix 25 g": "මස් කරි කුඩු 25 g",
    "meat curry mix 25gm": "මස් කරි කුඩු 25 g",
    "meat curry mix 250g": "මස් කරි කුඩු 250g",
    "meat curry mix 250gm": "මස් කරි කුඩු 250g",
    
    # Standardized lowercase keys to match structural text sanitization rules
    "w r string hopper flour 1kg": "සුදු සහල් ඉඳිආප්ප පිටි 1Kg",
    "r r string hopper flour 1kg": "රතු සහල් ඉඳිආප්ප පිටි 1Kg",
    "idliyas mixture 400gm": "ඉඩ්ලියාස් මිශ්‍රණය 400gm",
    "roasted r r string hopper flour 700gm": "බැදපු රතු සහල් ඉඳිආප්ප පිටි 700gm",
    
    "chilli powder 100gm": "මිරිස් කුඩු - 24x100gm",
    "chilli powder 50gm": "මිරිස් කුඩු - 40x50gm",
    "chilli powder 250gm": "මිරිස් කුඩු - 12x250gm",
    "chilli pieces 250gm": "කෑලි මිරිස් - 12x250gm",
    "chilli pieces 100gm": "කෑලි මිරිස් - 24x100gm",
    "chilli pieces 50gm": "කෑලි මිරිස් - 40x50gm",
    "curry powder 100gm": "තුනපහ කුඩු - 24x100gm",
    "curry powder 50gm": "තුනපහ කුඩු - 40x50gm",
    "curry powder 250gm": "තුනපහ කුඩු - 12x250gm",
    "r.curry powder 250gm": "බැදපු තුනපහ කුඩු - 12x250gm",
    "r.curry powder 100gm": "බැදපු තුනපහ කුඩු - 24x100gm",
    "r.curry powder 50gm": "බැදපු තුනපහ කුඩු - 40x50gm",
    "turmeric powder 100gm": "කහ කුඩු - 24x100gm",
    "turmeric powder 50gm": "කහ කුඩු - 40x50gm",
    "pepper powder 100gm": "ගම්මිරිස් කුඩු - 24x100gm",
    "pepper powder 50gm": "ගම්මිරිස් කුඩු - 40x50gm",
    "meat curry mix 100gm": "මස් කරි කුඩු - 24x100gm",
    "meat curry mix 50gm": "මස් කරි කුඩු - 40x50gm",
    "chilli powder 1 kg": "මිරිස් කුඩු 1 Kg",
    "chilli powder 1kg": "මිරිස් කුඩු 1 Kg",
    "chilli powder 500 g": "මිරිස් කුඩු 500 g",
    "chilli powder 500gm": "මිරිස් කුඩු 500 g",
    "roasted chilli powder 100 g": "බැදපු මිරිස් කුඩු 100 g",
    "roasted chilli powder 100gm": "බැදපු මිරිස් කුඩු 100 g",
    "chilli pieces 1 kg": "කෑලි මිරිස් 1 Kg",
    "chilli pieces 1kg": "කෑලි මිරිස් 1 Kg",
    "chilli pieces 500 g": "කෑලි මිරිස් 500 g",
    "chilli pieces 500gm": "කෑලි මිරිස් 500 g",
    "curry powder 1 kg": "තුනපහ කුඩු 1 Kg",
    "curry powder 1kg": "තුනපහ කුඩු 1 Kg",
    "pepper powder 500 g": "ගම්මිරිස් කුඩු 500 g",
    "pepper powder 500gm": "ගම්මිරිස් කුඩු 500 g",
    "white rice flour 1 kg": "සුදු සහල් පිටි 1 Kg",
    "white rice flour 1kg": "සුදු සහල් පිටි 1 Kg",
    "white rice flour 400 g": "සුදු සහල් පිටි 400 g",
    "white rice flour 400gm": "සුදු සහල් පිටි 400 g",
    "red rice flour 1 kg": "රතු සහල් පිටි 1 Kg",
    "red rice flour 1kg": "රතු සහල් පිටි 1 Kg",
    "red rice flour 400 g": "රතු සහල් පිටි 400 g",
    "red rice flour 400gm": "රතු සහල් පිටි 400 g",
    "undu flour 200 g": "උඳු පිටි 200 g",
    "undu flour 200gm": "උඳු පිටි 200 g",
    "kurakkan flour 400 g": "කුරක්කන් පිටි 400 g",
    "kurakkan flour 400g":"කුරක්කන් පිටි ග්‍රෑම් 400",
    "kurakkan flour 400gm": "කුරක්කන් පිටි 400 g",
    "kurakkan flour 200 gm": "කුරක්කන් පිටි 200 gm",
    "kurakkan flour 200g": "කුරක්කන් පිටි 200 gm",
    "gram flour 1kg": "කඩල පිටි 1Kg",
    "gram flour 1 kg": "කඩල පිටි 1Kg",
    "gram flour 200 g": "කඩල පිටි 200 g",
    "gram flour 200gm": "කඩල පිටි 200 g",
    "atta flour 400 g": "ආටා පිටි 400 g",
    "atta flour 400gm": "ආටා පිටි 400 g",
    "atta flour 1 kg": "ආටා පිටි 1 Kg",
    "atta flour 1kg": "ආටා පිටි 1 Kg",
    "white rice s/h flour 700 g": "සුදු සහල් පිටි 700 g",
    "white rice s/h flour 700gm": "සුදු සහල් පිටි 700 g",
    "red rice s/h flour 700 g": "රතු සහල් පිටි 700 g",
    "red rice s/h flour 700gm": "රතු සහල් පිටි 700 g",
    "hopper mixture 400 g": "ආප්ප මිශ්‍රණය 400 g",
    "hopper mixture 400gm": "ආප්ප මිශ්‍රණය 400 g",
    "thosei mixture 400 g": "තෝසේ මිශ්‍රණය 400 g",
    "thosei mixture 400gm": "තෝසේ මිශ්‍රණය 400 g",
    "idli mixture 400 g": "ඉඩ්ලි මිශ්‍රණය 400 g",
    "idli mixture 400gm": "ඉඩ්ලි මිශ්‍රණය 400 g",
    "soya promo pack": "සෝයා ප්‍රෝමෝ පැක්",
    "curry soya meat 50 g": "කරි සෝයා මීට් 50 g",
    "soya meat curry 50gm": "කරි සෝයා මීට් 50 g",
    "chicken soya meat 50 g": "චිකන් සෝයා මීට් 50 g",
    "soya meat chicken 50gm": "චිකන් සෝයා මීට් 50 g",
    "cuttle fish soya meat 50 g": "කට්ල්ෆිෂ් සෝයා මීට් 50 g",
    "table salt 400 g": "කුඩු ලුණු 400 g",
    "table salt 400gm": "කුඩු ලුණු 400 g",
    "cummin seed 100 g": "සූදුරු 100 g",
    "cumin seeds 100gm": "සූදුරු 100 g",
    "mustard seed 100 g": "අබ ඇට 100 g",
    "mustard seed 100gm": "අබ ඇට 100 g",
    "b/pepper seed 100 g": "කළු ගම්මිරිස් ඇට 100 g",
    "pepper seed 100gm": "කළු ගම්මිරිස් ඇට 100 g",
    "dill seed 100 g": "උළු හාල් 100 g",
    "dill seed 100gm": "උළු හාල් 100 g",
    "goraka 100 g": "ගොරකා 100 g",
    "goraka 100gm": "ගොරකා 100 g",
    "corriander 200 g": "කොත්තමල්ලි 200 g",
    "coriander 200gm": "කොත්තමල්ලි 200 g",
    "cinnamon stick 50 g": "කුරුඳු පොතු 50 g",
    "cinnamon stick 50gm": "කුරුඳු පොතු 50 g",
    "maldive fish(chips) 100 g": "උම්බලකඩ කෑලි 100 g",
    "maldive fish(chips) 100gm": "උම්බලකඩ කෑලි 100 g",
    "sago 200 g": "සව් 200 g",
    "sago 200gm": "සව් 200 g",
    "cloves 20 g": "කරාබුනැටි 20 g",
    "cloves 20gm": "කරාබුනැටි 20 g",
    "plums 100 g": "ප්ලම්ස් 100 g",
    "plums 100gm": "ප්ලම්ස් 100 g",
    "sukiri 100 g": "සූකිරි 100 g",
    "sukiri 100gm": "සූkිරි 100 g",
    "corriander 100 g": "කොත්තමල්ලි 100 g",
    "coriander seed 100gm": "කොත්තමල්ලි 100 g",
    "sago 100 g": "සව් 100 g",
    "sago 100gm": "සව් 100 g",
    "fennel seed 100 g": "මහදුරු 100 g",
    "fennel seeds 100gm": "මහදුරු 100 g",
    "maldive fish (chips) 50 g": "උම්බලකඩ කෑලි 50 g",
    "maldive fish (chips) 50gm": "උම්බලකඩ කෑලි 50 g",
    "maldive fish (chips) 25 g": "උම්බලකඩ කෑලි 25 g",
    "maldive fish (chips) 25gm": "උම්බලකඩ කෑලි 25 g",
    "cinnamon stick 20 g": "කුරුඳු පොතු 20 g",
    "cinnamon stick 20gm": "කුරුඳු පොතු 20 g",
    "umbalakada sambal 10 g": "උම්බලකඩ සම්බෝල 10 g",
    "umbalakada sambal 10gm": "උම්බලකඩ සම්බෝල 10 g",
    "rasam cream 100 g": "රසම් ක්‍රීම් 100 g",
    "rasam cream 100gm": "රසම් ක්‍රීම් 100 g",
    "mustard powder 50 g": "අබ කුඩු 50 g",
    "mustard powder 50gm": "අබ කුඩු 50 g",
    "goraka cream 100 g": "ගොරකා ක්‍රීම් 100 g",
    "goraka cream 100gm": "ගොරකා ක්‍රීම් 100 g",
    "tea 108x100gm": "තේ කොළ - 108x100gm",
    "tea 360x50gm": "තේ කොළ - 360x50gm",
    "coffee 100 g": "කෝපි 100 g",
    "coffee 100gm": "කෝපි 100 g",
    "coffee 50 g": "කෝපි 50 g",
    "coffee 50gm": "කෝපි 50 g",
    "coffee 20 g": "කෝපි 20 g",
    "coffee 20gm": "කෝපි 20 g",
    "papadam 50 g": "පපඩම් 50 g",
    "papadam 50gm": "පපඩම් 50 g",
    "cummin seed 50 g": "සූදුරු 50 g",
    "cumin seeds 50gm": "සූදුරු 50 g",
    "mustard seed 50 g": "අබ ඇට 50 g",
    "mustard seeds 50gm": "අබ ඇට 50 g",
    "b/pepper seed 50 g": "කළු ගම්මිරිස් ඇට 50 g",
    "pepper seeds 50gm": "කළු ගම්මිරිස් ඇට 50 g",
    "dill seed 50 g": "උළු හාල් 50 g",
    "dill seeds 50gm": "උළු හාල් 50 g",
    "fennel seed 50 g": "මහදුරු 50 g",
    "fennel seeds 50gm": "මහදුරු 50 g",
    "vinegar 750ml": "විනාකිරි 750ML",
    "vinegar 350ml": "විනාකිරි 350ML",
    "curry powder 500 g": "තුනපහ කුඩු 500 g",
    "curry powder 500gm": "තුනපහ කුඩු 500 g",
    "roasted curry powder 500 g": "බැදපු තුනපහ කුඩු 500 g",
    "roasted curry powder 500gm": "බැදපු තුනපහ කුඩු 500 g",
    "white rice s/h flour 1 kg": "සුදු සහල් පිටි 1 Kg",
    "white rice s/h flour 1kg": "සුදු සහල් පිටි 1 Kg",
    "red rice s/h flour 1 kg": "රතු සහල් පිටි 1 Kg",
    "red rice s/h flour 1kg": "රතු සහල් පිටි 1 Kg",
    "special noodles 400 g": "විශේෂ නූඩ්ල්ස් 400 g",
    "special noodles 400gm": "විශේෂ නූඩ්ල්ස් 400 g",
    "noodles 5 kg": "නූඩ්ල්ස් 5 Kg",
    "noodles 5kg": "නූඩ්ල්ස් 5 Kg",
    "table salt 1 kg": "ලුණු 1 Kg",
    "table salt 1kg": "ලුණු 1 Kg",
    "goraka 50 g": "ගොරකා 50 g",
    "goraka 50gm": "ගොරකා 50 g",
    "papadam 1 kg": "පපඩම් 1 kg",
    "papadam 1kg": "පපඩම් 1 kg",
    "kottu 100gm": "කොත්තු 100g"
}

def clean_text_for_matching(text):
    if not text:
        return ""
    cleaned = text.replace('"', '').replace('/', ' ').replace('.', ' ').replace('-', ' ')
    return " ".join(cleaned.lower().split())

def apply_column_widths(table):
    """Explicitly assigns structural widths across word table grid cells."""
    widths = [Inches(4.0), Inches(1.0), Inches(1.0), Inches(1.0)]
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

# UI Header Layout Elements
st.markdown('<div class="main-title">📋 රුහුණු භාණ්ඩ ලැයිස්තුව</div>', unsafe_allow_html=True)
st.markdown('<div class="sub-title">ඔබේ Picklist PDF එක සිංහල Word ගොනුවක් බවට ක්ෂණිකව පරිවර්තනය කරන්න</div>', unsafe_allow_html=True)

# Step 1 Container Instruction Banner
st.markdown('<div class="step-container"><strong>පියවර 1:</strong> ඔබේ මුල් පිටපතේ PDF ගොනුව පහත කොටුවට එක් කරන්න (Upload PDF File)</div>', unsafe_allow_html=True)

# File Uploader
uploaded_file = st.file_uploader("", type=["pdf"], label_visibility="collapsed")

if uploaded_file is not None:
    with st.spinner("දත්ත විශ්ලේෂණය කරමින් පවතී. කරුණාකර රැඳී සිටින්න..."):
        # Initialize structured Word Document
        doc = Document()
        doc.add_heading('රුහුණු භාණ්ඩ ලැයිස්තුව', level=1)
        
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # Define base layout headers
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'භාණ්ඩය'
        hdr_cells[1].text = 'නොමිලේ'
        hdr_cells[2].text = 'පෙට්ටි'
        hdr_cells[3].text = 'කොටස්'
        
        matched_count = 0
        preview_data = []
        
        with pdfplumber.open(uploaded_file) as pdf:
            for page in pdf.pages:
                text_content = page.extract_text()
                if not text_content:
                    continue
                
                lines = text_content.split('\n')
                for line in lines:
                    normalized_line = line.strip()
                    
                    if '|' in normalized_line:
                        parts = [p.strip() for p in normalized_line.split('|')]
                        
                        if len(parts) >= 7:
                            desc = parts[1]
                            matchable_desc = clean_text_for_matching(desc)
                            
                            for english_key, sinhala_val in PRODUCT_MAPPING.items():
                                matchable_key = clean_text_for_matching(english_key)
                                
                                if matchable_key == matchable_desc or matchable_key in matchable_desc:
                                    try:
                                        free_qty = parts[4]
                                        cs_qty = parts[5]
                                        pcs_qty = parts[6]
                                        
                                        row_cells = table.add_row().cells
                                        row_cells[0].text = sinhala_val
                                        row_cells[1].text = free_qty
                                        row_cells[2].text = cs_qty
                                        row_cells[3].text = pcs_qty
                                        
                                        preview_data.append({
                                            "භාණ්ඩය": sinhala_val, 
                                            "නොමිලේ": free_qty, 
                                            "පෙට්ටි": cs_qty, 
                                            "කොටස්": pcs_qty
                                        })
                                        matched_count += 1
                                        break
                                    except IndexError:
                                        continue
                    else:
                        matchable_line = clean_text_for_matching(normalized_line)
                        for english_key, sinhala_val in PRODUCT_MAPPING.items():
                            matchable_key = clean_text_for_matching(english_key)
                            
                            if matchable_key in matchable_line:
                                all_numbers = re.findall(r'\b\d+(?:\.\d+)?\b', normalized_line)
                                if len(all_numbers) >= 4:
                                    free_qty = all_numbers[-4]
                                    cs_qty = all_numbers[-3]
                                    pcs_qty = all_numbers[-2]
                                    
                                    row_cells = table.add_row().cells
                                    row_cells[0].text = sinhala_val
                                    row_cells[1].text = free_qty
                                    row_cells[2].text = cs_qty
                                    row_cells[3].text = pcs_qty
                                    
                                    preview_data.append({
                                        "භාණ්ඩය": sinhala_val, 
                                        "නොමිලේ": free_qty, 
                                        "පෙට්ටි": cs_qty, 
                                        "කොටස්": pcs_qty
                                    })
                                    matched_count += 1
                                    break
        
        # Enforce explicitly customized column widths in Word document output
        apply_column_widths(table)

    if matched_count > 0:
        st.success(f"🎉 සාර්ථකයි! ගැළපෙන භාණ්ඩ පේළි {matched_count} ක් සාර්ථකව පරිවර්තනය කරන ලදී.")
        
        # Step 2 Container Instruction Banner
        st.markdown('<div class="step-container"><strong>පියවර 2:</strong> සකස් කරන ලද නව දත්ත පෙරදසුන පරීක්ෂා කර බාගත කරගන්න</div>', unsafe_allow_html=True)
        
        # Displaying preview inside the custom layout card
        st.markdown('<div class="preview-card">', unsafe_allow_html=True)
        st.dataframe(preview_data, use_container_width=True)
        st.markdown('</div>', unsafe_allow_html=True)
        
        st.write("")  # Spacer Element
        
        # Render output document to byte storage array stream
        doc_stream = io.BytesIO()
        doc.save(doc_stream)
        doc_stream.seek(0)
        
        # Action download button
        st.download_button(
            label="📥 නිපදවන ලද Word ලිපිගොනුව බාගත කරගන්න (Download Document)",
            data=doc_stream,
            file_name="රුහුණු_භාණ්ඩ_ලැයිස්තුව.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
    else:
        st.error("⚠️ දෝෂයකි: අප්ලෝඩ් කරන ලද PDF ගොනුවේ තිබූ කිසිදු භාණ්ඩයක් අපගේ නාමාවලිය සමඟ ගැළපුණේ නැත. කරුණාකර වෙනත් ගොනුවක් උත්සාහ කරන්න.")
